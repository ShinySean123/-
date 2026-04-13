import streamlit as st
import pandas as pd
import re
import math
import io
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, Border, Side
from openpyxl.utils import get_column_letter

# Word 處理相關
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 設定網頁標題與寬度
st.set_page_config(page_title="題庫格式轉換神器", page_icon="🚀", layout="centered")

st.title("🚀 題庫格式轉換神器")
st.markdown("上傳原始 Excel，一秒自動排版出完美的 **Word 試卷** 與 **Excel 精修版**！")

# ==================== UI 介面 ====================
with st.container():
    st.subheader("1. 上傳題庫")
    uploaded_file = st.file_uploader("請上傳原始題庫檔案 (.xlsx)", type=["xlsx"])

with st.container():
    st.subheader("2. 設定輸出檔名")
    col1, col2 = st.columns(2)
    with col1:
        exam_title = st.text_input("考卷標題 (Word 大標題)", "113年 測驗題庫")
    with col2:
        excel_filename = st.text_input("Excel 輸出檔名", "精修版題庫")

# 清理檔名函數
def sanitize(name): 
    return re.sub(r'[\\/:*?"<>|]', '_', str(name))

# ==================== 核心處理邏輯 ====================
if uploaded_file is not None:
    if st.button("✨ 開始轉換 ✨", use_container_width=True):
        with st.spinner("系統正在施展魔法中，請稍候..."):
            try:
                # --- 1. 智慧讀取 Excel ---
                df_raw = pd.read_excel(uploaded_file, header=None)
                header_idx = -1
                for i, row in df_raw.iterrows():
                    row_str = " ".join([str(x) for x in row if pd.notna(x)])
                    if any(k in row_str for k in ["題目", "Question", "內容"]):
                        header_idx = i
                        break
                
                # 重新以正確的 header 讀取
                uploaded_file.seek(0) # 重置指標
                df = pd.read_excel(uploaded_file, header=header_idx if header_idx != -1 else 0)

                # --- 2. 欄位偵測 ---
                df.columns = [str(c).strip() for c in df.columns]
                cols = df.columns.tolist()

                q_col = next((c for c in cols if any(k in c for k in ["題目", "內容", "Question"])), None)
                ans_col = next((c for c in cols if any(k in c for k in ["答案", "Answer", "Correct", "判定"])), None)
                expl_col = next((c for c in cols if any(k in c for k in ["詳解", "說明", "Explanation"])), None)
                src_col = next((c for c in cols if any(k in c for k in ["出處", "來源", "Source"])), None)
                
                # 動態抓取選項
                opt_cols = [c for c in cols if ("選項" in c or "Option" in c or re.match(r'^[A-F]$', c)) 
                            and c not in [q_col, ans_col, expl_col, src_col]]
                
                if len(opt_cols) < 2:
                    opt_cols = [c for c in cols if c in ['A', 'B', 'C', 'D', 'E', 'F'] and c not in [q_col, ans_col, expl_col, src_col]]

                opt_cols.sort()
                max_opts = len(opt_cols)
                opt_labels = [chr(65 + i) for i in range(max_opts)]

                processed_rows = []
                def clean_text(text):
                    return str(text).replace('$', '').strip() if pd.notna(text) else ""

                # --- 3. 資料整理 ---
                for i, row in df.iterrows():
                    q_text = clean_text(row.get(q_col, ""))
                    if not q_text: continue
                    
                    row_dict = {'題號': len(processed_rows) + 1, '題目內容': q_text}
                    for idx, label in enumerate(opt_labels):
                        row_dict[f'選項{label}'] = clean_text(row.get(opt_cols[idx], ""))
                    
                    ans_raw = clean_text(row.get(ans_col, ""))
                    match = re.search(r'[A-F]', ans_raw.upper())
                    row_dict['正確答案'] = match.group(0) if match else ""
                    row_dict['針對各選項之詳解'] = clean_text(row.get(expl_col, ""))
                    row_dict['出處'] = clean_text(row.get(src_col, ""))
                    processed_rows.append(row_dict)

                # ==================== 產出 Excel (記憶體緩衝) ====================
                excel_buffer = io.BytesIO()
                pd.DataFrame(processed_rows).to_excel(excel_buffer, index=False)
                excel_buffer.seek(0)
                
                wb = load_workbook(excel_buffer)
                ws = wb.active
                
                ans_idx = 3 + max_opts
                col_widths = {'A': 8, 'B': 40}
                for i in range(max_opts): col_widths[get_column_letter(3+i)] = 30
                col_widths[get_column_letter(ans_idx)] = 15
                col_widths[get_column_letter(ans_idx+1)] = 60
                col_widths[get_column_letter(ans_idx+2)] = 40

                for letter, width in col_widths.items(): ws.column_dimensions[letter].width = width
                border = Border(top=Side(style='thin'), bottom=Side(style='thin'), left=Side(style='thin'), right=Side(style='thin'))
                
                for r_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=ws.max_row), 1):
                    max_lines = 1
                    for cell in row:
                        cell.border = border
                        cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                        if r_idx == 1:
                            cell.font = Font(bold=True)
                            cell.alignment = Alignment(horizontal='center', vertical='center')
                        if cell.column_letter in ['A', get_column_letter(ans_idx)]:
                            cell.alignment = Alignment(horizontal='center', vertical='center')
                        
                        if r_idx > 1:
                            c_w = col_widths.get(cell.column_letter, 20)
                            est = math.ceil((len(str(cell.value)) * 1.8) / c_w)
                            if est > max_lines: max_lines = est
                    if r_idx > 1: ws.row_dimensions[r_idx].height = max_lines * 18
                
                # 將美化後的 Excel 存入新的 Buffer
                final_excel_buffer = io.BytesIO()
                wb.save(final_excel_buffer)
                final_excel_buffer.seek(0)

                # ==================== 產出 Word (記憶體緩衝) ====================
                word_buffer = io.BytesIO()
                doc = Document()
                sec = doc.sections[0]
                sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(1.27)
                doc.styles['Normal'].font.name = 'Times New Roman'
                doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
                doc.styles['Normal'].font.size = Pt(12)
                
                PURPLE, BLUE = RGBColor(112, 48, 160), RGBColor(0, 50, 150)

                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(exam_title)
                run.bold = True
                run.font.size = Pt(16)

                for row in processed_rows:
                    doc.add_paragraph(f"{row['题號']}. {row['題目內容']}").paragraph_format.space_after = Pt(6)
                    for lbl in opt_labels:
                        txt = row.get(f'選項{lbl}', '')
                        if txt:
                            opt_p = doc.add_paragraph(f"({lbl}) {txt}")
                            opt_p.paragraph_format.left_indent, opt_p.paragraph_format.space_after = Pt(18), Pt(0)
                    
                    ans_p = doc.add_paragraph()
                    ans_p.paragraph_format.space_before = Pt(6)
                    ans_p.add_run("Ans : ").bold = True
                    ans_p.add_run(f"({row['正確答案']})")

                    expl = str(row['針對各選項之詳解'])
                    if expl and expl != "nan":
                        h = doc.add_paragraph()
                        h.paragraph_format.space_before, h.paragraph_format.space_after = Pt(4), Pt(0)
                        r = h.add_run("詳解 :"); r.bold, r.font.color.rgb = True, PURPLE
                        for line in expl.split('\n'):
                            if not line.strip(): continue
                            lp = doc.add_paragraph(); lp.paragraph_format.left_indent, lp.paragraph_format.space_after = Pt(18), Pt(2)
                            m = re.match(r'^([A-F])\s*([\(（].*?[\)）]|[:：])', line.strip())
                            if m:
                                pre, rest = m.group(0), line.strip()[len(m.group(0)):]
                                r1 = lp.add_run(pre); r1.bold, r1.font.color.rgb = True, PURPLE
                                r2 = lp.add_run(rest); r2.font.color.rgb = PURPLE
                            else:
                                lp.add_run(line.strip()).font.color.rgb = PURPLE
                    
                    src = str(row['出處'])
                    if src and src != "nan":
                        sp = doc.add_paragraph(); sp.paragraph_format.space_before, sp.paragraph_format.space_after = Pt(2), Pt(0)
                        r_l = sp.add_run("出處 : "); r_l.bold, r_l.font.color.rgb = True, BLUE
                        sp.add_run(src).font.color.rgb = BLUE
                    doc.add_paragraph("")
                
                doc.save(word_buffer)
                word_buffer.seek(0)

                st.success("🎉 轉換成功！請點擊下方按鈕下載檔案。")
                
                # --- 顯示下載按鈕 ---
                st.download_button(
                    label="📄 下載 Word 考卷",
                    data=word_buffer,
                    file_name=f"{sanitize(exam_title)}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.download_button(
                    label="📊 下載 Excel 精修版",
                    data=final_excel_buffer,
                    file_name=f"{sanitize(excel_filename)}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

            except Exception as e:
                st.error(f"發生錯誤：{e}")